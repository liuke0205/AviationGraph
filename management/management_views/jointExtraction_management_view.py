import os
import re
import threading
import docx
from django.contrib import messages
from django.http import JsonResponse, FileResponse
from django.shortcuts import render, redirect

from AviationGraph.ChinsesNER.main import ChineseNER
# 跳转到关系查询页面
from AviationGraph.utils.MySqlConn import MySqlConn
from AviationGraph.utils.utils import write_excel_xls


def toJointExtraction(request):
    return render(request, 'management/joint_extraction.html')


def jointExtraction_upload(request):
    if request.method == 'POST':
        # 获取文件名
        from builtins import str
        path = request.FILES.get('file')
        str_filename = str(path)

        if str_filename.endswith(".docx") or str_filename.endswith(".doc"):
            if path:
                if os.path.isfile("upload_file/joint_extraction_word.docx"):
                    os.remove("upload_file/joint_extraction_word.docx")
                with open('upload_file/joint_extraction_word.docx', 'wb+') as destination:
                    for chunk in path.chunks():
                        destination.write(chunk)
                '''
                先将所有语句取出换行和\t 去除，然后合到一起，最后再按照句号分句
                '''
                new_data = []
                for p in docx.Document("upload_file/joint_extraction_word.docx").paragraphs:
                    data = p.text.replace('\n', '').replace('\t', '')
                    new_data.append(data)
                all_word = "".join(new_data).replace(" ", "")
                print(all_word)
                request.session['all_word'] = all_word
                return render(request, 'management/joint_extraction.html', {'str': all_word})
            else:
                messages.success(request, "文件为空！")
                return redirect('/management/toJointExtraction/')
        else:
            messages.success(request, "上传失败，请上传Word文件！")
            return redirect('/management/toJointExtraction/')

# 将抽取结果写入到excel文件的线程
class myThread (threading.Thread):
    def __init__(self, threadID, name, resultList):
        threading.Thread.__init__(self)
        self.threadID = threadID
        self.name = name
        self.resultList = resultList
    def run(self):
        print("开始线程，去写入文件：" + self.name)
        write_excel_xls("download_file/jointExact_admin.xls", "jointExact", self.resultList)
        print("退出线程，结束写文件：" + self.name)

from AviationGraph.utils.snowFlake import IdWorker

# 读取待识别文本并转化成所要形式
def joint_extraction(request):
    # 1.先读取upload_file/relation_extraction_word.docx 看是否存在
    import os.path
    if os.path.isfile("upload_file/joint_extraction_word.docx"):
        # 3.如果已经上传进行抽取，抽取完后删除
        new_data = []
        for p in docx.Document("upload_file/joint_extraction_word.docx").paragraphs:
            data = p.text.replace('\n', '').replace('\t', '')
            new_data.append(data)
        all_word = "".join(new_data).replace(" ", "")

        pattern = r'。|！|；|;'
        senetence_list = re.split(pattern, str(all_word))
        new_senetence_list = []
        for data in senetence_list:
            if len(data) > 0:
                new_senetence_list.append(data + "。")

        # 将session域初始化
        request.session['resultList'] = []

        resultList = []
        cn = ChineseNER("predict")
        global num_progress
        for i in range(len(new_senetence_list)):
            # 调用深度学习预测模型
            temp_list = cn.predict(new_senetence_list[i])
            for data in temp_list:
                if len(data) > 0:
                    snowker = IdWorker(1, 1)
                    id = snowker.get_id()
                    print(type(data), data)
                    data.append(id)
                    resultList.append(data)
            num_progress = round(i * 100 / len(new_senetence_list), 2)
        # 抽取完毕后，删除word文件
        if os.path.isfile("upload_file/joint_extraction_word.docx"):
            os.remove("upload_file/joint_extraction_word.docx")

        if os.path.isfile("download_file/jointExact_admin.xls"):
            os.remove("download_file/jointExact_admin.xls")

        # 开启一个线程 将结果保存到xlsx文件中
        thread1 = myThread(1, "Thread-1", resultList)
        thread1.start()

        request.session['resultList3'] = resultList
        return redirect('/management/display_re_text/')
    # 2.如果不存在提示还没上传文件
    else:
        messages.success(request, "请上传Word文件！")
        return redirect('/management/toJointExtraction/')


def show_progress(request):
    return JsonResponse(num_progress, safe=False)


def download_jointExtract_result(request):
    if os.path.isfile("download_file/jointExact_admin.xls"):
        file = open('download_file/jointExact_admin.xls', 'rb')
        response = FileResponse(file)
        response['Content-Type'] = 'application/octet-stream'
        response['Content-Disposition'] = 'attachment;filename="jointExact_admin.xls"'
        return response
    else:
        messages.success(request, "请先进行联合抽取！")
        return redirect('/management/toJointExtraction/')


def display_re_text(request):
    resultList = request.session.get('resultList3')
    if resultList is None or len(resultList) == 0:
        messages.success(request, "请先进行联合抽取！")
        return redirect('/management/toJointExtraction/')
    else:
        return render(request, 'management/joint_extracting_result.html', {'resultList': resultList})

def addCorpus(request):
    resultList = request.session.get('resultList3')
    id = request.GET.get('id')
    print(id)
    conn = MySqlConn('source').connectMySql()
    for idx in range(len(resultList)):
        data = resultList[idx]
        if str(data[4]) == str(id):
            cursor = conn.cursor()
            insert_sql = "INSERT INTO corpus(head_entity, relation, tail_entity, text) VALUES ('%s','%s','%s','%s')" % \
                         (data[0], data[1], data[2], data[3])
            cursor.execute(insert_sql)
            conn.commit()
            del resultList[idx]
            break
    request.session['resultList3'] = resultList
    return redirect('/management/display_re_text/')


def deleteCorpus(request):
    resultList = request.session.get('resultList3')
    id = request.GET.get('id')
    for idx in range(len(resultList)):
        data = resultList[idx]
        if str(data[4]) == str(id):
            del resultList[idx]
            break
    request.session['resultList3'] = resultList
    return redirect('/management/display_re_text/')

def modifyCorpus(request):
    resultList = request.session.get('resultList3')
    id = request.POST.get('rel_id')
    headEntity = request.POST.get('headEntity')
    relation = request.POST.get('relation')
    tailEntity = request.POST.get('tailEntity')
    text = request.POST.get('text')
    for idx in range(len(resultList)):
        if str(resultList[idx][4]) == str(id):
            print(resultList[idx])
            resultList[idx][0] = headEntity
            resultList[idx][1] = relation
            resultList[idx][2] = tailEntity
            resultList[idx][3] = text
            break
    request.session['resultList3'] = resultList
    return redirect('/management/display_re_text/')
